import hashlib
import math
import json
import csv
from pathlib import Path
from typing import List, Dict, Any, Optional
from io import BytesIO
import re

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import pypdf
import docx

from models import (
    Document,
    DocumentChunk,
    DocumentClassification,
    RAGQuery,
    RAGResult,
    CollectionName,
    RetrievalResult,
)


class RAGService:
    """
    RAG Service with ChromaDB (persistent) — multi-collection architecture.

    Collections:
      - compliance_documents: heading-aware chunked compliance/regulatory docs
      - stakeholder_requirements: atomic (unchunked) stakeholder requirements
      - user_stories: atomic (unchunked) generated user stories

    IMPORTANT:
    - semantic_search() returns top_k candidates ALWAYS (no min_similarity filtering)
    - thresholding is done at answer layer
    """

    HEADING_RE = re.compile(r"^(\d+(\.\d+)*)\s+.+|^(SECTION|CHAPTER)\s+\w+|^[A-Z][A-Z\s\-]{6,}$")

    def __init__(self):
        print("[INFO] Loading embedding model...")
        self.embedding_model = SentenceTransformer("BAAI/bge-small-en-v1.5", device="cpu")
        print("[OK] Embedding model loaded")

        chroma_path = Path("storage/chroma")
        chroma_path.mkdir(parents=True, exist_ok=True)

        self.chroma_client = chromadb.PersistentClient(
            path=str(chroma_path),
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True,
            ),
        )

        # Three collections with cosine distance
        self.compliance_collection = self.chroma_client.get_or_create_collection(
            name="compliance_documents",
            metadata={"hnsw:space": "cosine"},
        )
        self.requirements_collection = self.chroma_client.get_or_create_collection(
            name="stakeholder_requirements",
            metadata={"hnsw:space": "cosine"},
        )
        self.stories_collection = self.chroma_client.get_or_create_collection(
            name="user_stories",
            metadata={"hnsw:space": "cosine"},
        )

        # Backward compat alias — existing code that references self.collection
        self.collection = self.compliance_collection

        # Clean break: delete legacy single collection if it exists
        try:
            self.chroma_client.delete_collection("reqpal_documents")
            print("[DEL] Deleted legacy reqpal_documents collection")
        except Exception:
            pass

        print(f"[DB] Collections: compliance={self.compliance_collection.count()}, "
              f"requirements={self.requirements_collection.count()}, "
              f"stories={self.stories_collection.count()}")

        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)

    # =========================================================
    # Document processing (compliance documents)
    # =========================================================

    async def process_document(
            self,
            file_content: bytes,
            filename: str,
            project_id: int,
            classification: DocumentClassification,
            document_id: int,
            metadata: Dict[str, Any] = {},
    ) -> Document:
        file_ext = Path(filename).suffix.lower().replace(".", "")
        file_hash = hashlib.md5(file_content).hexdigest()
        safe_name = f"{project_id}_{file_hash}_{filename}"
        file_path = self.upload_dir / safe_name

        with open(file_path, "wb") as f:
            f.write(file_content)

        document = Document(
            id=document_id,
            project_id=project_id,
            filename=filename,
            file_type=file_ext,
            file_path=str(file_path),
            file_size=len(file_content),
            classification=classification,
            processing_status="processing",
            metadata=metadata,
        )

        text = await self._extract_text(file_content, file_ext)
        chunks = await self._chunk_text(text, document.id, filename)

        await self._embed_and_store(chunks, project_id, document)

        document.chunks = chunks
        document.total_chunks = len(chunks)
        document.indexed = True
        document.processing_status = "completed"

        return document

    async def _extract_text(self, content: bytes, file_type: str) -> str:
        if file_type == "pdf":
            reader = pypdf.PdfReader(BytesIO(content))
            pages = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(f"[Page {i+1}]\n{t}")
            return "\n\n".join(pages)

        if file_type == "docx":
            d = docx.Document(BytesIO(content))
            return "\n\n".join(p.text for p in d.paragraphs if (p.text or "").strip())

        if file_type == "txt":
            return content.decode("utf-8", errors="ignore")

        if file_type == "csv":
            txt = content.decode("utf-8", errors="ignore")
            return "\n".join(" | ".join(r) for r in csv.reader(txt.splitlines()))

        if file_type == "json":
            return json.dumps(json.loads(content.decode("utf-8")), indent=2)

        raise ValueError(f"Unsupported file type: {file_type}")

    def _is_heading(self, line: str) -> bool:
        line = (line or "").strip()
        return bool(self.HEADING_RE.match(line))

    async def _chunk_text(
        self,
        text: str,
        document_id: int,
        filename: str,
        chunk_size: int = 1300,
        overlap: int = 240,
    ) -> List[DocumentChunk]:
        """
        Heading-aware chunking:
        1) Build blocks by headings + paragraph-ish boundaries
        2) Pack blocks into size-limited chunks with char-overlap
        """

        lines = [l.rstrip() for l in text.splitlines()]
        blocks: List[tuple[str | None, str]] = []
        buf: List[str] = []
        current_heading: str | None = None

        def flush_buf():
            nonlocal buf
            if buf:
                blocks.append((current_heading, "\n".join(buf).strip()))
                buf = []

        for line in lines:
            if not (line or "").strip():
                continue

            if self._is_heading(line):
                flush_buf()
                current_heading = line.strip()
                continue

            buf.append(line)

            if line.endswith(".") and len(" ".join(buf)) > 280:
                flush_buf()

        flush_buf()

        chunks: List[DocumentChunk] = []
        chunk_index = 0
        cur = ""
        cur_heading: str | None = None

        def add_chunk(content: str, heading: str | None):
            nonlocal chunk_index
            chunk_id = f"doc{document_id}_chunk{chunk_index}"
            chunks.append(
                DocumentChunk(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    content=content.strip(),
                    chunk_index=chunk_index,
                    section=heading,
                    metadata={"filename": filename, "total_length": len(content)},
                )
            )
            chunk_index += 1

        for heading, block in blocks:
            block_text = f"{heading}\n{block}" if heading else block

            if len(cur) + len(block_text) + 2 > chunk_size and cur.strip():
                add_chunk(cur, cur_heading)

                tail = cur[-overlap:] if len(cur) > overlap else cur
                cur = (tail + "\n\n" + block_text).strip()
                cur_heading = heading or cur_heading
            else:
                if not cur_heading and heading:
                    cur_heading = heading
                cur = (cur + "\n\n" + block_text).strip() if cur else block_text

        if cur.strip():
            add_chunk(cur, cur_heading)

        return chunks

    async def _embed_and_store(
        self,
        chunks: List[DocumentChunk],
        project_id: int,
        document: Document,
        collection: Optional[chromadb.Collection] = None,
    ):
        if not chunks:
            return

        target = collection or self.compliance_collection

        texts = [c.content for c in chunks]
        ids = [c.chunk_id for c in chunks]

        embeddings = self.embedding_model.encode(
            texts,
            show_progress_bar=False,
            normalize_embeddings=True,
        )

        metadatas = [
            {
                "project_id": int(project_id),
                "document_id": int(document.id),
                "document_name": document.filename,
                "classification": document.classification.value,
                "chunk_index": int(c.chunk_index),
                "filename": c.metadata.get("filename"),
            }
            for c in chunks
        ]

        target.add(
            ids=ids,
            documents=texts,
            embeddings=embeddings.tolist(),
            metadatas=metadatas,
        )

        for c in chunks:
            c.embedding_id = c.chunk_id

        print(f"[STORE] Indexed {len(chunks)} chunks in {target.name}")

    # =========================================================
    # Delete helpers
    # =========================================================

    def delete_document_chunks(self, document_id: int) -> int:
        """Delete all chunks for a document from the compliance collection."""
        try:
            existing = self.compliance_collection.get(
                where={"document_id": int(document_id)}
            )
            count = len(existing["ids"]) if existing.get("ids") else 0
            if count > 0:
                self.compliance_collection.delete(
                    where={"document_id": int(document_id)}
                )
            print(f"[DEL] Deleted {count} chunks for document {document_id}")
            return count
        except Exception as e:
            print(f"Warning: delete_document_chunks failed: {e}")
            return 0

    def delete_stakeholder_requirement(self, project_id: int, requirement_id: int) -> None:
        emb_id = f"req_{project_id}_{requirement_id}"
        try:
            self.requirements_collection.delete(ids=[emb_id])
        except Exception as e:
            print(f"Warning: delete requirement embedding failed: {e}")

    def delete_user_story(self, project_id: int, story_id: int) -> None:
        emb_id = f"story_{project_id}_{story_id}"
        try:
            self.stories_collection.delete(ids=[emb_id])
        except Exception as e:
            print(f"Warning: delete story embedding failed: {e}")

    def delete_project_requirements(self, project_id: int) -> None:
        """Delete all stakeholder requirements for a project from the vector store."""
        try:
            self.requirements_collection.delete(
                where={"project_id": int(project_id)}
            )
        except Exception as e:
            print(f"Warning: delete project requirements failed: {e}")

    def delete_project_stories(self, project_id: int) -> None:
        """Delete all user stories for a project from the vector store."""
        try:
            self.stories_collection.delete(
                where={"project_id": int(project_id)}
            )
        except Exception as e:
            print(f"Warning: delete project stories failed: {e}")

    # =========================================================
    # Stakeholder requirement indexing (atomic, no chunking)
    # =========================================================

    async def add_stakeholder_requirement(
        self,
        project_id: int,
        requirement_id: int,
        title: str,
        description: str,
        stakeholder_role: str,
        category: str = "functional",
        priority: str = "Must Have",
        source: str = "manual",
        import_batch_id: Optional[str] = None,
    ) -> str:
        """Embed a stakeholder requirement as a single atomic document. NO chunking."""
        text = f"{title}\n{description}"

        embedding = self.embedding_model.encode(
            [text],
            normalize_embeddings=True,
        )

        emb_id = f"req_{project_id}_{requirement_id}"

        chroma_metadata: Dict[str, Any] = {
            "project_id": int(project_id),
            "requirement_id": int(requirement_id),
            "title": title,
            "stakeholder_role": stakeholder_role,
            "category": category,
            "priority": priority,
            "source": source,
        }
        if import_batch_id:
            chroma_metadata["import_batch_id"] = import_batch_id

        self.requirements_collection.upsert(
            ids=[emb_id],
            documents=[text],
            embeddings=embedding.tolist(),
            metadatas=[chroma_metadata],
        )

        print(f"[STORE] Indexed stakeholder requirement {emb_id}")
        return emb_id

    # =========================================================
    # User story indexing (atomic, no chunking)
    # =========================================================

    async def add_user_story(
        self,
        project_id: int,
        story_id: int,
        title: str,
        description: str,
        acceptance_criteria: str = "",
        requirement_id: Optional[int] = None,
        category: str = "functional",
    ) -> str:
        """Embed a user story as a single atomic document. NO chunking."""
        parts = [title, description]
        if acceptance_criteria:
            parts.append(f"Acceptance Criteria: {acceptance_criteria}")
        text = "\n".join(parts)

        embedding = self.embedding_model.encode(
            [text],
            normalize_embeddings=True,
        )

        emb_id = f"story_{project_id}_{story_id}"

        chroma_metadata: Dict[str, Any] = {
            "project_id": int(project_id),
            "story_id": int(story_id),
            "title": title,
            "category": category,
        }
        if requirement_id is not None:
            chroma_metadata["requirement_id"] = int(requirement_id)

        self.stories_collection.upsert(
            ids=[emb_id],
            documents=[text],
            embeddings=embedding.tolist(),
            metadatas=[chroma_metadata],
        )

        print(f"[STORE] Indexed user story {emb_id}")
        return emb_id

    # =========================================================
    # Semantic search (legacy — compliance only, no thresholding)
    # =========================================================

    async def semantic_search(self, query: RAGQuery) -> List[RAGResult]:
        print(f"[SEARCH] RAG search: '{query.query}'")

        q_emb = self.embedding_model.encode(
            [query.query],
            normalize_embeddings=True,
        )[0]

        where_filter: Dict[str, Any] = {"project_id": query.project_id}
        if query.document_classifications:
            where_filter["classification"] = {"$in": [c.value for c in query.document_classifications]}

        results = self.compliance_collection.query(
            query_embeddings=[q_emb.tolist()],
            n_results=query.top_k,
            where=where_filter,
        )

        ids = results["ids"][0] if results.get("ids") else []
        metadatas = results["metadatas"][0] if results.get("metadatas") else []
        docs = results["documents"][0] if results.get("documents") else []
        distances = results["distances"][0] if results.get("distances") else []

        out: List[RAGResult] = []

        for i, chunk_id in enumerate(ids):
            md = metadatas[i] or {}
            content = docs[i] or ""
            dist = float(distances[i]) if i < len(distances) else 0.0
            sim = 1.0 - dist

            out.append(
                RAGResult(
                    chunk_id=chunk_id,
                    document_id=int(md.get("document_id") or 0),
                    document_name=str(md.get("document_name", "Unknown")),
                    content=content,
                    similarity_score=float(sim),
                    classification=str(md.get("classification", "unknown")),
                    page_number=md.get("page_number"),
                    section=md.get("section") or md.get("section"),
                )
            )

        return out

    # =========================================================
    # Multi-collection retrieval
    # =========================================================

    async def retrieve_context(
        self,
        query: str,
        project_id: int,
        retrieval_mode: str = "all",
        document_classifications: Optional[List[DocumentClassification]] = None,
        top_k: int = 10,
    ) -> List[RetrievalResult]:
        """
        Multi-collection retrieval.
        - "compliance": only compliance_documents collection
        - "requirements": only stakeholder_requirements collection
        - "stories": only user_stories collection
        - "all": query all three, merge, sort by similarity
        """
        q_emb = self.embedding_model.encode(
            [query], normalize_embeddings=True
        )[0].tolist()

        results: List[RetrievalResult] = []

        if retrieval_mode in ("all", "compliance"):
            results.extend(
                self._query_compliance(q_emb, project_id, document_classifications or [], top_k)
            )

        if retrieval_mode in ("all", "requirements"):
            results.extend(
                self._query_requirements(q_emb, project_id, top_k)
            )

        if retrieval_mode in ("all", "stories"):
            results.extend(
                self._query_stories(q_emb, project_id, top_k)
            )

        results.sort(key=lambda r: r.similarity_score, reverse=True)
        return results[:top_k]

    def _query_compliance(
        self,
        q_emb: list,
        project_id: int,
        classifications: List[DocumentClassification],
        top_k: int,
    ) -> List[RetrievalResult]:
        if classifications:
            where_filter = {
                "$and": [
                    {"project_id": int(project_id)},
                    {"classification": {"$in": [c.value for c in classifications]}},
                ]
            }
        else:
            where_filter = {"project_id": int(project_id)}

        try:
            raw = self.compliance_collection.query(
                query_embeddings=[q_emb],
                n_results=top_k,
                where=where_filter,
            )
        except Exception as e:
            print(f"Warning: compliance query failed: {e}")
            return []

        return self._parse_compliance_results(raw)

    def _query_requirements(
        self,
        q_emb: list,
        project_id: int,
        top_k: int,
    ) -> List[RetrievalResult]:
        try:
            raw = self.requirements_collection.query(
                query_embeddings=[q_emb],
                n_results=top_k,
                where={"project_id": int(project_id)},
            )
        except Exception as e:
            print(f"Warning: requirements query failed: {e}")
            return []

        return self._parse_requirements_results(raw)

    def _query_stories(
        self,
        q_emb: list,
        project_id: int,
        top_k: int,
    ) -> List[RetrievalResult]:
        try:
            raw = self.stories_collection.query(
                query_embeddings=[q_emb],
                n_results=top_k,
                where={"project_id": int(project_id)},
            )
        except Exception as e:
            print(f"Warning: stories query failed: {e}")
            return []

        return self._parse_stories_results(raw)

    def _parse_compliance_results(self, raw: dict) -> List[RetrievalResult]:
        ids = raw["ids"][0] if raw.get("ids") else []
        metadatas = raw["metadatas"][0] if raw.get("metadatas") else []
        docs = raw["documents"][0] if raw.get("documents") else []
        distances = raw["distances"][0] if raw.get("distances") else []

        out = []
        for i, chunk_id in enumerate(ids):
            md = metadatas[i] or {}
            sim = 1.0 - float(distances[i]) if i < len(distances) else 0.0
            out.append(RetrievalResult(
                chunk_id=chunk_id,
                document_id=int(md.get("document_id", 0)),
                document_name=str(md.get("document_name", "Unknown")),
                content=docs[i] or "",
                similarity_score=sim,
                source_collection=CollectionName.COMPLIANCE,
                classification=md.get("classification"),
                page_number=md.get("page_number"),
                section=md.get("section"),
            ))
        return out

    def _parse_requirements_results(self, raw: dict) -> List[RetrievalResult]:
        ids = raw["ids"][0] if raw.get("ids") else []
        metadatas = raw["metadatas"][0] if raw.get("metadatas") else []
        docs = raw["documents"][0] if raw.get("documents") else []
        distances = raw["distances"][0] if raw.get("distances") else []

        out = []
        for i, emb_id in enumerate(ids):
            md = metadatas[i] or {}
            sim = 1.0 - float(distances[i]) if i < len(distances) else 0.0
            out.append(RetrievalResult(
                chunk_id=emb_id,
                requirement_id=int(md.get("requirement_id", 0)),
                document_name=md.get("title", "Stakeholder Requirement"),
                content=docs[i] or "",
                similarity_score=sim,
                source_collection=CollectionName.REQUIREMENTS,
                stakeholder_role=md.get("stakeholder_role"),
                category=md.get("category"),
            ))
        return out

    def _parse_stories_results(self, raw: dict) -> List[RetrievalResult]:
        ids = raw["ids"][0] if raw.get("ids") else []
        metadatas = raw["metadatas"][0] if raw.get("metadatas") else []
        docs = raw["documents"][0] if raw.get("documents") else []
        distances = raw["distances"][0] if raw.get("distances") else []

        out = []
        for i, emb_id in enumerate(ids):
            md = metadatas[i] or {}
            sim = 1.0 - float(distances[i]) if i < len(distances) else 0.0
            out.append(RetrievalResult(
                chunk_id=emb_id,
                story_id=int(md.get("story_id", 0)),
                document_name=md.get("title", "User Story"),
                content=docs[i] or "",
                similarity_score=sim,
                source_collection=CollectionName.USER_STORIES,
                category=md.get("category"),
            ))
        return out

    # =========================================================
    # Two-stage retrieval with reranker
    # =========================================================

    async def retrieve_and_rerank(
        self,
        query: str,
        project_id: int,
        retrieval_mode: str = "all",
        document_classifications: Optional[List[DocumentClassification]] = None,
        top_k: int = 8,
    ) -> List[RetrievalResult]:
        """
        Two-stage retrieval:
          Stage 1: Get top 20 candidates via embedding search
          Stage 2: Rerank with CrossEncoder, select top_k
        """
        from backend.services.reranker_service import reranker_service, RerankItem

        # Stage 1: broad retrieval
        candidates = await self.retrieve_context(
            query=query,
            project_id=project_id,
            retrieval_mode=retrieval_mode,
            document_classifications=document_classifications,
            top_k=20,
        )

        if not candidates:
            return []

        # Stage 2: rerank with cross-encoder
        rerank_items = [
            RerankItem(
                chunk_id=c.chunk_id,
                document_name=c.document_name,
                classification=c.classification or c.category or "unknown",
                content=c.content,
                base_similarity=c.similarity_score,
            )
            for c in candidates
        ]

        reranked = reranker_service.rerank(
            question=query,
            candidates=rerank_items,
            top_n=top_k,
        )

        # Map back to RetrievalResult with reranker score
        result_map = {c.chunk_id: c for c in candidates}
        out = []
        for item, score in reranked:
            original = result_map.get(item.chunk_id)
            if original:
                original.similarity_score = 1.0 / (1.0 + math.exp(-score))
                out.append(original)

        return out

    # =========================================================
    # Convenience: search similar requirements only
    # =========================================================

    async def search_similar_requirements(
        self,
        query: str,
        project_id: int,
        top_k: int = 5,
    ) -> List[RetrievalResult]:
        """Search only the stakeholder_requirements collection."""
        return await self.retrieve_context(
            query=query,
            project_id=project_id,
            retrieval_mode="requirements",
            top_k=top_k,
        )

    # =========================================================
    # Stats
    # =========================================================

    def get_stats(self) -> Dict[str, Any]:
        return {
            "vector_db": "ChromaDB",
            "collections": {
                "compliance_documents": self.compliance_collection.count(),
                "stakeholder_requirements": self.requirements_collection.count(),
                "user_stories": self.stories_collection.count(),
            },
            "embedding_model": "BAAI/bge-small-en-v1.5",
            "space": "cosine",
        }


rag_service = RAGService()
__all__ = ["rag_service", "RAGService"]
